import streamlit as st
import os
import zipfile
import re
from html import escape
from dotenv import load_dotenv

import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# -----------------------------
# Vercel Serverless Entrypoint Compatibility
# -----------------------------
class handler(BaseHTTPRequestHandler):
    def do_GET(self):
        self.send_response(200)
        self.send_header('Content-type', 'text/html')
        self.end_headers()
        html_content = """<!DOCTYPE html>
<html>
<head>
    <title>AI Documentation Generator</title>
    <style>
        body { font-family: system-ui, -apple-system, sans-serif; display: flex; justify-content: center; align-items: center; min-height: 100vh; margin: 0; background: #0f172a; color: white; }
        .card { background: #1e293b; padding: 2.5rem; border-radius: 12px; box-shadow: 0 10px 25px rgba(0,0,0,0.5); text-align: center; max-width: 520px; border: 1px solid #334155; }
        h1 { margin-top: 0; color: #38bdf8; font-size: 1.8rem; }
        p { line-height: 1.6; color: #94a3b8; margin-bottom: 1.5rem; }
        .badge { background: #0369a1; color: #e0f2fe; padding: 4px 10px; border-radius: 6px; font-size: 0.85rem; font-weight: bold; }
    </style>
</head>
<body>
    <div class="card">
        <h1>📄 AI Documentation Generator</h1>
        <p><span class="badge">Vercel Deployment Active</span></p>
        <p>This Streamlit app is configured for deployment. Run locally using <code>streamlit run app.py</code> or host on Streamlit Cloud for full interactive UI.</p>
    </div>
</body>
</html>"""
        self.wfile.write(html_content.encode('utf-8'))

app = handler
application = handler

# -----------------------------
# Gemini Configuration & Helpers
# -----------------------------
load_dotenv()

@st.cache_data(ttl=300)
def fetch_gemini_models(key):
    """Fetch available models supporting generateContent from the Gemini API."""
    fallback_models = ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-1.5-flash", "gemini-2.5-pro"]
    if not key:
        return fallback_models
    try:
        genai.configure(api_key=key)
        models = []
        for m in genai.list_models():
            if "generateContent" in m.supported_generation_methods:
                name = m.name.replace("models/", "")
                models.append(name)
        if models:
            return models
    except Exception:
        pass
    return fallback_models

# -----------------------------
# Create Output Folder
# -----------------------------
os.makedirs("output", exist_ok=True)

# -----------------------------
# Streamlit Page
# -----------------------------
st.set_page_config(
    page_title="AI Documentation Generator",
    page_icon="📄",
    layout="wide"
)

st.title("📄 AI Documentation Generator")
st.write("Upload your source code file and generate professional AI documentation.")

# -----------------------------
# Sidebar Configuration & Model Selection
# -----------------------------
with st.sidebar:
    st.header("🔑 API Key Settings")
    user_api_key = st.text_input(
        "Gemini API Key",
        type="password",
        help="Enter your own Gemini API key or leave empty to use default server key.",
        placeholder="AIzaSy..."
    )

    env_api_key = os.getenv("GEMINI_API_KEY")
    active_api_key = user_api_key.strip() if user_api_key else env_api_key

    if active_api_key:
        genai.configure(api_key=active_api_key)
        if user_api_key:
            st.success("Using Visitor API Key", icon="🔑")
        else:
            st.info("Using Default Server API Key", icon="🌐")
    else:
        st.warning("⚠️ No Gemini API Key set. Enter your key above.")

    st.markdown("---")
    st.header("⚙️ Model Settings")
    
    available_models = fetch_gemini_models(active_api_key)
    
    default_index = available_models.index("gemini-2.5-flash") if "gemini-2.5-flash" in available_models else 0
    selected_model_name = st.selectbox(
        "Select Gemini Model",
        options=available_models,
        index=default_index,
        help="Models are dynamically fetched from your Gemini API account."
    )
    st.caption(f"Active Model: `{selected_model_name}`")

    st.markdown("---")
    st.header("📘 About")
    st.write("**AI Documentation Generator**")
    st.write("Generate documentation automatically from source code using AI.")

    st.markdown("---")
    st.subheader("Features")
    st.write("✅ Upload source code")
    st.write("✅ AI documentation")
    st.write("✅ Dynamic Model Selection")
    st.write("✅ Formatted PDF Export")
    st.write("✅ Formatted DOCX Export")
    st.write("✅ Markdown Export")
    st.write("✅ ZIP Download")

    st.markdown("---")
    st.info("Version 1.2")

# -----------------------------
# Helper: Docx Inline Markdown Parser
# -----------------------------
def add_docx_markdown_run(paragraph, text):
    """Parses basic markdown inline tags (**bold**, *italic*, `code`) into docx runs."""
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    tokens = pattern.split(text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(token)

# -----------------------------
# Save DOCX (Formatted)
# -----------------------------
def save_docx(text):
    filename = "output/documentation.docx"
    document = Document()

    document.add_heading("AI Generated Documentation", level=1)

    lines = text.split("\n")
    in_code_block = False
    code_lines = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
        elif re.match(r'^[\*\-\+]\s+', stripped):
            content = re.sub(r'^[\*\-\+]\s+', '', stripped)
            p = document.add_paragraph(style='List Bullet')
            add_docx_markdown_run(p, content)
        elif re.match(r'^\d+\.\s+', stripped):
            content = re.sub(r'^\d+\.\s+', '', stripped)
            p = document.add_paragraph(style='List Number')
            add_docx_markdown_run(p, content)
        else:
            p = document.add_paragraph()
            add_docx_markdown_run(p, stripped)

    document.save(filename)
    return filename

# -----------------------------
# Helper: PDF HTML Parser & Sanitizer
# -----------------------------
def md_to_pdf_html(line):
    # Replace common unicode quotes/dashes and sanitize for ReportLab Latin-1 Helvetica font
    line = line.replace('’', "'").replace('“', '"').replace('”', '"').replace('–', '-').replace('—', '-')
    line = line.encode('latin-1', 'ignore').decode('latin-1')
    line = escape(line)
    line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
    line = re.sub(r'\*(.*?)\*', r'<i>\1</i>', line)
    line = re.sub(r'`(.*?)`', r'<font face="Courier" color="#c7254e">\1</font>', line)
    return line

# -----------------------------
# Save PDF (Formatted Platypus Layout)
# -----------------------------
def save_pdf(text):
    filename = "output/documentation.pdf"
    doc = SimpleDocTemplate(
        filename,
        pagesize=letter,
        rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Title'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#1E293B'),
        spaceAfter=15
    )

    h1_style = ParagraphStyle(
        'H1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=15,
        leading=19,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'H2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=17,
        textColor=colors.HexColor('#2563EB'),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )

    h3_style = ParagraphStyle(
        'H3',
        parent=styles['Heading3'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=15,
        textColor=colors.HexColor('#334155'),
        spaceBefore=8,
        spaceAfter=4,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#334155'),
        spaceAfter=6
    )

    bullet_style = ParagraphStyle(
        'Bullet',
        parent=body_style,
        leftIndent=15,
        spaceAfter=3
    )

    code_style = ParagraphStyle(
        'Code',
        fontName='Courier',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#0F172A'),
        backColor=colors.HexColor('#F1F5F9'),
        borderColor=colors.HexColor('#E2E8F0'),
        borderWidth=1,
        borderPadding=6,
        spaceBefore=6,
        spaceAfter=8
    )

    story = [Paragraph("AI Generated Documentation", title_style), Spacer(1, 10)]

    lines = text.split("\n")
    in_code_block = False
    code_lines = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                code_text = escape("\n".join(code_lines))
                story.append(Preformatted(code_text, code_style))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            continue

        if stripped.startswith("# "):
            story.append(Paragraph(md_to_pdf_html(stripped[2:]), h1_style))
        elif stripped.startswith("## "):
            story.append(Paragraph(md_to_pdf_html(stripped[3:]), h2_style))
        elif stripped.startswith("### "):
            story.append(Paragraph(md_to_pdf_html(stripped[4:]), h3_style))
        elif re.match(r'^[\*\-\+]\s+', stripped):
            content = re.sub(r'^[\*\-\+]\s+', '', stripped)
            story.append(Paragraph(f"• {md_to_pdf_html(content)}", bullet_style))
        elif re.match(r'^\d+\.\s+', stripped):
            story.append(Paragraph(md_to_pdf_html(stripped), bullet_style))
        else:
            story.append(Paragraph(md_to_pdf_html(stripped), body_style))

    doc.build(story)
    return filename

# -----------------------------
# Save Markdown
# -----------------------------
def save_markdown(text):
    filename = "output/documentation.md"
    with open(filename, "w", encoding="utf-8") as file:
        file.write(text)
    return filename

# -----------------------------
# Create ZIP
# -----------------------------
def create_zip():
    zip_name = "output/documentation.zip"
    with zipfile.ZipFile(zip_name, "w") as zipf:
        files = [
            "output/documentation.pdf",
            "output/documentation.docx",
            "output/documentation.md"
        ]
        for file in files:
            if os.path.exists(file):
                zipf.write(
                    file,
                    arcname=os.path.basename(file)
                )
    return zip_name

# -----------------------------
# File Upload & Generation
# -----------------------------
uploaded_file = st.file_uploader(
    "Choose a source code file",
    type=["py", "java", "cpp", "c", "js", "html", "css", "txt"]
)

if uploaded_file is not None:

    # Read uploaded file
    code = uploaded_file.read().decode("utf-8")

    st.success("✅ File uploaded successfully!")

    st.subheader("📜 Source Code")

    st.code(code)

    if st.button("🤖 Generate AI Documentation"):
        if not active_api_key:
            st.error("❌ No Gemini API Key provided. Please enter your Gemini API key in the sidebar.")
        else:
            try:
                genai.configure(api_key=active_api_key)
                with st.spinner(f"Generating documentation using `{selected_model_name}`..."):
                    model = genai.GenerativeModel(selected_model_name)
                    
                    prompt = f"""
You are an expert software documentation engineer.

Analyze the following source code and generate detailed professional documentation.

Include the following sections:

1. Project Overview
2. Purpose
3. Features
4. Technologies Used
5. Function Description
6. Class Description
7. Inputs and Outputs
8. Installation Guide
9. Example Usage
10. Conclusion

Source Code:

{code}
"""

                    response = model.generate_content(prompt)
                    documentation = response.text
                    st.success("✅ Documentation Generated Successfully!")

                st.subheader("📄 Generated Documentation")

                st.markdown(documentation)

                # -----------------------------
                # Generate Files
                # -----------------------------
                docx_file = save_docx(documentation)

                pdf_file = save_pdf(documentation)

                md_file = save_markdown(documentation)

                zip_file = create_zip()
                # -----------------------------
                # Download Buttons
                # -----------------------------
                st.markdown("---")
                st.subheader("📥 Download Documentation")

                col1, col2, col3, col4 = st.columns(4)

                with col1:
                    with open(docx_file, "rb") as f:
                        st.download_button(
                            label="📄 DOCX",
                            data=f,
                            file_name="documentation.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                with col2:
                    with open(pdf_file, "rb") as f:
                        st.download_button(
                            label="📕 PDF",
                            data=f,
                            file_name="documentation.pdf",
                            mime="application/pdf"
                        )

                with col3:
                    with open(md_file, "rb") as f:
                        st.download_button(
                            label="📝 Markdown",
                            data=f,
                            file_name="documentation.md",
                            mime="text/markdown"
                        )

                with col4:
                    with open(zip_file, "rb") as f:
                        st.download_button(
                            label="📦 ZIP",
                            data=f,
                            file_name="documentation.zip",
                            mime="application/zip"
                        )

            except Exception as e:
                st.error(f"❌ Error: {e}")

# -----------------------------
# Footer
# -----------------------------
st.markdown("---")

st.caption("Developed by Dani Toffin | AI Documentation Generator | 2026")
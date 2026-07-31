from flask import Flask, request, jsonify, render_template, send_file
import os
import io
import zipfile
import re
import uuid
from html import escape
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

import zlib
import base64
import tempfile
import hashlib

load_dotenv()

template_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'templates'))
if not os.path.exists(template_dir):
    template_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), 'templates'))

app = Flask(__name__, template_folder=template_dir)

# Temporary in-memory cache for generated documentation downloads
DOC_CACHE = {}

# In-memory authentication database & active sessions
USERS_DB = {
    "demo@neurodocs.ai": {
        "name": "Dani Toffin",
        "email": "demo@neurodocs.ai",
        "password_hash": hashlib.sha256("password123".encode()).hexdigest(),
        "created_at": "2026-07-30"
    }
}
USER_SESSIONS = {}

def encode_doc_token(doc_text):
    try:
        compressed = zlib.compress(doc_text.encode('utf-8'))
        return base64.urlsafe_b64encode(compressed).decode('ascii')
    except Exception:
        return str(uuid.uuid4())

def decode_doc_token(token):
    if not token:
        return None
    if token in DOC_CACHE:
        return DOC_CACHE[token]
    tmp_path = os.path.join(tempfile.gettempdir(), f"neurodoc_{token[:64]}.txt")
    if os.path.exists(tmp_path):
        try:
            with open(tmp_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception:
            pass
    try:
        compressed = base64.urlsafe_b64decode(token.encode('ascii'))
        return zlib.decompress(compressed).decode('utf-8')
    except Exception:
        pass
    return None

def get_active_api_key(custom_key=None):
    if custom_key and custom_key.strip():
        return custom_key.strip()
    return os.getenv("GEMINI_API_KEY", "").strip()

def fetch_gemini_models(api_key):
    fallback_models = ["gemini-2.0-flash", "gemini-2.5-flash", "gemini-flash-latest", "gemini-2.0-flash-lite", "gemini-pro-latest"]
    if not api_key:
        return fallback_models
    try:
        genai.configure(api_key=api_key)
        raw_models = []
        for m in genai.list_models():
            if "generateContent" in m.supported_generation_methods:
                name = m.name.replace("models/", "")
                if not any(bad in name for bad in ["antigravity", "lyria", "robotics", "computer-use", "tts"]):
                    raw_models.append(name)
        ordered = [m for m in fallback_models if m in raw_models]
        others = [m for m in raw_models if m not in fallback_models]
        result = ordered + others
        if result:
            return result
    except Exception:
        pass
    return fallback_models

def add_docx_markdown_run(paragraph, text):
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    tokens = pattern.split(text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(token)

def generate_docx_bytes(text):
    document = Document()
    document.add_heading("AI Generated Documentation", level=1)
    lines = text.split("\n")
    in_code_block = False
    code_lines = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code_block:
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
        elif re.match(r'^[\*\-\+]\s+', stripped):
            content = re.sub(r'^[\*\-\+]\s+', '', stripped)
            p = document.add_paragraph(style='List Bullet')
            add_docx_markdown_run(p, content)
        elif re.match(r'^\d+\.\s+', stripped):
            content = re.sub(r'^\d+\.\s+', '', stripped)
            p = document.add_paragraph(style='List Number')
            add_docx_markdown_run(p, content)
        else:
            p = document.add_paragraph()
            add_docx_markdown_run(p, stripped)

    if in_code_block and code_lines:
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)

    bio = io.BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio

def md_to_pdf_html(line):
    line = line.replace('’', "'").replace('“', '"').replace('”', '"').replace('–', '-').replace('—', '-')
    line = line.encode('latin-1', 'ignore').decode('latin-1')
    line = escape(line)
    line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
    line = re.sub(r'\*(.*?)\*', r'<i>\1</i>', line)
    line = re.sub(r'`(.*?)`', r'<font face="Courier" color="#c7254e">\1</font>', line)
    return line

def safe_pdf_paragraph(text, style):
    try:
        return Paragraph(text, style)
    except Exception:
        clean = escape(re.sub(r'<[^>]+>', '', text))
        return Paragraph(clean, style)

def generate_pdf_bytes(text):
    bio = io.BytesIO()
    doc = SimpleDocTemplate(
        bio,
        pagesize=letter,
        rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Title'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#1E293B'),
        spaceAfter=15
    )
    h1_style = ParagraphStyle(
        'H1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=15,
        leading=19,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )
    h2_style = ParagraphStyle(
        'H2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=17,
        textColor=colors.HexColor('#2563EB'),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )
    h3_style = ParagraphStyle(
        'H3',
        parent=styles['Heading3'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=15,
        textColor=colors.HexColor('#334155'),
        spaceBefore=8,
        spaceAfter=4,
        keepWithNext=True
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#334155'),
        spaceAfter=6
    )
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=body_style,
        leftIndent=15,
        spaceAfter=3
    )
    code_style = ParagraphStyle(
        'Code',
        fontName='Courier',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#0F172A'),
        backColor=colors.HexColor('#F1F5F9'),
        borderColor=colors.HexColor('#E2E8F0'),
        borderWidth=1,
        borderPadding=6,
        spaceBefore=6,
        spaceAfter=8
    )

    story = [safe_pdf_paragraph("AI Generated Documentation", title_style), Spacer(1, 10)]
    lines = text.split("\n")
    in_code_block = False
    code_lines = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code_block:
                code_text = escape("\n".join(code_lines))
                story.append(Preformatted(code_text, code_style))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            continue

        if stripped.startswith("# "):
            story.append(safe_pdf_paragraph(md_to_pdf_html(stripped[2:]), h1_style))
        elif stripped.startswith("## "):
            story.append(safe_pdf_paragraph(md_to_pdf_html(stripped[3:]), h2_style))
        elif stripped.startswith("### "):
            story.append(safe_pdf_paragraph(md_to_pdf_html(stripped[4:]), h3_style))
        elif re.match(r'^[\*\-\+]\s+', stripped):
            content = re.sub(r'^[\*\-\+]\s+', '', stripped)
            story.append(safe_pdf_paragraph(f"• {md_to_pdf_html(content)}", bullet_style))
        elif re.match(r'^\d+\.\s+', stripped):
            story.append(safe_pdf_paragraph(md_to_pdf_html(stripped), bullet_style))
        else:
            story.append(safe_pdf_paragraph(md_to_pdf_html(stripped), body_style))

    if in_code_block and code_lines:
        code_text = escape("\n".join(code_lines))
        story.append(Preformatted(code_text, code_style))

    doc.build(story)
    bio.seek(0)
    return bio

@app.route('/api/auth/signup', methods=['POST'])
def auth_signup():
    data = request.get_json(silent=True) or {}
    email = data.get('email', '').strip().lower()
    password = data.get('password', '')
    name = data.get('name', '').strip() or email.split('@')[0].capitalize()

    if not email or '@' not in email:
        return jsonify({"success": False, "message": "Please enter a valid email address."}), 400
    if not password or len(password) < 6:
        return jsonify({"success": False, "message": "Password must be at least 6 characters long."}), 400

    if email in USERS_DB:
        return jsonify({"success": False, "message": "An account with this email already exists."}), 400

    pwd_hash = hashlib.sha256(password.encode()).hexdigest()
    USERS_DB[email] = {
        "name": name,
        "email": email,
        "password_hash": pwd_hash,
        "created_at": "2026-07-30"
    }

    session_token = f"neuro_{uuid.uuid4().hex}"
    USER_SESSIONS[session_token] = email

    return jsonify({
        "success": True,
        "message": f"Welcome to NeuroDocs AI, {name}!",
        "token": session_token,
        "user": {"name": name, "email": email}
    })

@app.route('/api/auth/login', methods=['POST'])
def auth_login():
    data = request.get_json(silent=True) or {}
    email = data.get('email', '').strip().lower()
    password = data.get('password', '')

    if not email or not password:
        return jsonify({"success": False, "message": "Email and password are required."}), 400

    user = USERS_DB.get(email)
    if not user:
        return jsonify({"success": False, "message": "Invalid email or password."}), 401

    pwd_hash = hashlib.sha256(password.encode()).hexdigest()
    if user["password_hash"] != pwd_hash:
        return jsonify({"success": False, "message": "Invalid email or password."}), 401

    session_token = f"neuro_{uuid.uuid4().hex}"
    USER_SESSIONS[session_token] = email

    return jsonify({
        "success": True,
        "message": f"Welcome back, {user['name']}!",
        "token": session_token,
        "user": {"name": user["name"], "email": user["email"]}
    })

@app.route('/api/auth/me', methods=['POST'])
def auth_me():
    data = request.get_json(silent=True) or {}
    token = data.get('token', '')
    email = USER_SESSIONS.get(token)

    if not email or email not in USERS_DB:
        return jsonify({"authenticated": False})

    user = USERS_DB[email]
    return jsonify({
        "authenticated": True,
        "user": {"name": user["name"], "email": user["email"]}
    })

@app.route('/api/auth/logout', methods=['POST'])
def auth_logout():
    data = request.get_json(silent=True) or {}
    token = data.get('token', '')
    USER_SESSIONS.pop(token, None)
    return jsonify({"success": True, "message": "Logged out successfully."})

@app.route('/api/validate-key', methods=['POST'])
def validate_key():
    data = request.get_json(silent=True) or {}
    user_key = data.get('api_key', '').strip()
    active_key = get_active_api_key(user_key)
    
    if not active_key:
        return jsonify({
            "valid": False,
            "message": "No API key configured. Enter a key or configure server GEMINI_API_KEY."
        }), 400
    
    try:
        genai.configure(api_key=active_key)
        models = []
        for m in genai.list_models():
            if "generateContent" in m.supported_generation_methods:
                models.append(m.name.replace("models/", ""))
        return jsonify({
            "valid": True,
            "is_custom": bool(user_key),
            "model_count": len(models),
            "message": f"API Key Validated! {len(models)} models available."
        })
    except Exception as e:
        return jsonify({
            "valid": False,
            "message": f"API Key Validation Error: {str(e)}"
        }), 400

@app.route('/api/models', methods=['POST'])
def get_models():
    data = request.get_json(silent=True) or {}
    user_key = data.get('api_key', '')
    active_key = get_active_api_key(user_key)
    models = fetch_gemini_models(active_key)
    return jsonify({
        "models": models,
        "has_key": bool(active_key),
        "is_custom": bool(user_key and user_key.strip())
    })

def try_generate_content(model_name, prompt, api_key):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    response = model.generate_content(prompt)
    try:
        if response.text and response.text.strip():
            return response.text
    except Exception:
        pass
    if hasattr(response, 'candidates') and response.candidates:
        parts = response.candidates[0].content.parts
        txt = "\n".join([p.text for p in parts if hasattr(p, 'text')])
        if txt.strip():
            return txt
    raise ValueError(f"Model '{model_name}' did not return valid text content.")

@app.route('/api/generate', methods=['POST'])
def generate_docs():
    data = request.get_json(silent=True) or {}
    code = data.get('code', '')
    model_name = data.get('model', 'gemini-2.0-flash')
    user_key = data.get('api_key', '')
    sections = data.get('sections', [])

    if not code or not code.strip():
        return jsonify({"error": "Please enter or upload source code first."}), 400

    active_key = get_active_api_key(user_key)
    if not active_key:
        return jsonify({"error": "No Gemini API Key configured. Click 'Key Settings' to enter a key."}), 400

    if sections and isinstance(sections, list) and len(sections) > 0:
        sections_str = "\n".join([f"{idx+1}. {s}" for idx, s in enumerate(sections)])
    else:
        sections_str = """1. Project Overview & Purpose
2. Key Features & Capabilities
3. Architecture & Structure
4. Function & Method Descriptions
5. Class & Data Models
6. Inputs, Outputs & API Schema
7. Installation & Setup Guide
8. Example Usage & Quickstart"""

    prompt = f"""
You are an expert software documentation engineer.

Analyze the following source code and generate detailed professional documentation.

Include the following specific sections:
{sections_str}

Source Code:

{code}
"""

    fallback_candidates = [model_name, "gemini-2.5-flash", "gemini-flash-latest", "gemini-2.0-flash-lite", "gemini-1.5-flash", "gemini-pro-latest"]
    models_to_try = []
    for m in fallback_candidates:
        if m and m not in models_to_try:
            models_to_try.append(m)

    documentation = None
    used_model = model_name
    last_exception = None

    for m in models_to_try:
        try:
            documentation = try_generate_content(m, prompt, active_key)
            used_model = m
            break
        except Exception as e:
            last_exception = e
            continue

    if not documentation:
        err_msg = str(last_exception) if last_exception else "Generation failed"
        if "ResourceExhausted" in err_msg or "429" in err_msg or "quota" in err_msg.lower():
            return jsonify({
                "error": "Gemini API free quota limit reached for this minute. Please wait 30 seconds or click 'Key Settings' in top bar to use your own free key."
            }), 429
        return jsonify({"error": err_msg}), 500

    doc_id = encode_doc_token(documentation)
    DOC_CACHE[doc_id] = documentation

    try:
        tmp_path = os.path.join(tempfile.gettempdir(), f"neurodoc_{doc_id[:64]}.txt")
        with open(tmp_path, 'w', encoding='utf-8') as f:
            f.write(documentation)
    except Exception:
        pass

    return jsonify({
        "success": True,
        "token": doc_id,
        "model_used": used_model,
        "is_fallback": (used_model != model_name),
        "documentation": documentation
    })

@app.route('/api/download/<format_type>', methods=['GET', 'POST'])
def download_file(format_type):
    token = request.args.get('token', '')
    if not token and request.is_json:
        token = (request.get_json(silent=True) or {}).get('token', '')

    documentation = decode_doc_token(token)

    if not documentation:
        return "Invalid or expired download token.", 404

    if format_type == 'docx':
        bio = generate_docx_bytes(documentation)
        return send_file(
            bio,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="documentation.docx"
        )
    elif format_type == 'pdf':
        bio = generate_pdf_bytes(documentation)
        return send_file(
            bio,
            mimetype="application/pdf",
            as_attachment=True,
            download_name="documentation.pdf"
        )
    elif format_type == 'md':
        bio = io.BytesIO(documentation.encode('utf-8'))
        bio.seek(0)
        return send_file(
            bio,
            mimetype="text/markdown",
            as_attachment=True,
            download_name="documentation.md"
        )
    elif format_type == 'zip':
        docx_bio = generate_docx_bytes(documentation)
        pdf_bio = generate_pdf_bytes(documentation)

        zip_bio = io.BytesIO()
        with zipfile.ZipFile(zip_bio, 'w', zipfile.ZIP_DEFLATED) as zipf:
            zipf.writestr("documentation.md", documentation.encode('utf-8'))
            zipf.writestr("documentation.docx", docx_bio.getvalue())
            zipf.writestr("documentation.pdf", pdf_bio.getvalue())
        zip_bio.seek(0)

        return send_file(
            zip_bio,
            mimetype="application/zip",
            as_attachment=True,
            download_name="documentation.zip"
        )

    return "Invalid format specified.", 400

@app.route('/')
def home():
    return render_template('index.html')

if __name__ == '__main__':
    try:
        from waitress import serve
        print("Starting Production WSGI Server on http://127.0.0.1:5000 ...")
        serve(app, host='0.0.0.0', port=5000)
    except ImportError:
        app.run(host='0.0.0.0', port=5000, debug=True)
